#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
國語科期末複習卷 Word & Markdown 產出腳本
解析 HTML 題目內容並輸出符合 108 課綱技能規格的實體 Word 練習卷與 Markdown 說明檔。
"""

import os
import re
import html
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

def set_font(run, font_name='Microsoft JhengHei', size_pt=10.5, bold=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def add_formatted_text(paragraph, text, font_name='Microsoft JhengHei', size_pt=10.5, bold=False, color_rgb=None):
    parts = text.split('**')
    for idx, part in enumerate(parts):
        is_bold_part = (idx % 2 == 1)
        run = paragraph.add_run(part)
        set_font(run, font_name, size_pt, bold or is_bold_part, color_rgb)

def strip_html_tags(text):
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = html.unescape(text)
    # Clean up double spacing and leading/trailing spaces on each line
    lines = [line.strip() for line in text.split('\n')]
    return '\n'.join(line for line in lines if line)

def parse_html_exam(html_path):
    with open(html_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 1. Parse Reading Passage
    passage_match = re.search(r'<div class="reading-passage">([\s\S]*?)<\/div>', content)
    passage = {}
    if passage_match:
        passage_html = passage_match.group(1)
        title_match = re.search(r'<h4>(.*?)</h4>', passage_html)
        title = title_match.group(1) if title_match else "閱讀短文"
        paragraphs = re.findall(r'<p>(.*?)</p>', passage_html)
        passage['title'] = strip_html_tags(title)
        passage['paragraphs'] = [strip_html_tags(p) for p in paragraphs]
    
    # 2. Parse Section Titles
    sections = re.findall(r'<h3 class="section-title">\s*<span>(.*?)</span>\s*</h3>', content)
    
    # 3. Parse Questions
    question_cards = re.findall(r'<div class="question-card" data-qindex="(\d+)" data-correct="([A-D])">([\s\S]*?)<\/div>\s*(?:<!-- Question|\s*<\/div>\s*<!-- ====================)', content)
    # If not matched, try simpler match
    if not question_cards:
        question_cards = re.findall(r'<div class="question-card" data-qindex="(\d+)" data-correct="([A-D])">([\s\S]*?)<\/div>\s*<\/div>', content)
    
    questions = []
    for qindex, correct, qcontent in question_cards:
        qtext_match = re.search(r'<div class="question-text">([\s\S]*?)</div>', qcontent)
        qtext = strip_html_tags(qtext_match.group(1)) if qtext_match else ""
        
        options = re.findall(r'<span class="option-text">\s*(.*?)\s*</span>', qcontent)
        options = [strip_html_tags(opt) for opt in options]
        
        explanation_content_match = re.search(r'<div class="explanation-content">([\s\S]*?)</div>\s*</div>', qcontent)
        if explanation_content_match:
            exp_html = explanation_content_match.group(1)
            badges = re.findall(r'<span class="explanation-badge">【(.*?)】</span>', exp_html)
            exp_divs = re.findall(r'<div>([\s\S]*?)</div>', exp_html)
            
            exp_details = []
            for div in exp_divs:
                exp_details.append(strip_html_tags(div))
            explanation = "\n".join(exp_details)
        else:
            explanation = ""
            
        questions.append({
            'qindex': int(qindex),
            'correct': correct,
            'text': qtext,
            'options': options,
            'explanation': explanation
        })
        
    return passage, sections, questions

def build_docx(passage, sections, questions, output_path):
    doc = Document()
    
    # Page Setup - Margins
    sections_list = doc.sections
    for section in sections_list:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('四年級下學期國語科期末段考複習卷（第7-12課）')
    set_font(run_title, font_name='Microsoft JhengHei', size_pt=16, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run('班級：_____________   座號：______   姓名：_______________   得分：___________')
    set_font(run_sub, font_name='Microsoft JhengHei', size_pt=11, bold=True)
    p_sub.paragraph_format.space_after = Pt(24)
    
    # Questions by Sections
    sec_idx = 0
    for q_idx, q in enumerate(questions):
        # Insert Section Header if needed
        # We have 15 questions: Q1-2: Sec 0, Q3-4: Sec 1, Q5-7: Sec 2, Q8-12: Sec 3, Q13-15: Sec 4
        if q_idx == 0 and len(sections) > 0:
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(18)
            p_sec.paragraph_format.space_after = Pt(8)
            run = p_sec.add_run(sections[0])
            set_font(run, font_name='Microsoft JhengHei', size_pt=13, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
        elif q_idx == 2 and len(sections) > 1:
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(18)
            p_sec.paragraph_format.space_after = Pt(8)
            run = p_sec.add_run(sections[1])
            set_font(run, font_name='Microsoft JhengHei', size_pt=13, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
        elif q_idx == 4 and len(sections) > 2:
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(18)
            p_sec.paragraph_format.space_after = Pt(8)
            run = p_sec.add_run(sections[2])
            set_font(run, font_name='Microsoft JhengHei', size_pt=13, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
        elif q_idx == 7 and len(sections) > 3:
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(18)
            p_sec.paragraph_format.space_after = Pt(8)
            run = p_sec.add_run(sections[3])
            set_font(run, font_name='Microsoft JhengHei', size_pt=13, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
        elif q_idx == 12:
            # Insert Reading Passage before Section 5
            if len(sections) > 4:
                p_sec = doc.add_paragraph()
                p_sec.paragraph_format.space_before = Pt(18)
                p_sec.paragraph_format.space_after = Pt(8)
                run = p_sec.add_run(sections[4])
                set_font(run, font_name='Microsoft JhengHei', size_pt=13, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
            
            # Reading box
            p_box_start = doc.add_paragraph()
            p_box_start.paragraph_format.space_before = Pt(10)
            run = p_box_start.add_run("-------------------- 閱讀短文 --------------------")
            set_font(run, font_name='Microsoft JhengHei', size_pt=11, bold=True, color_rgb=RGBColor(0x7F, 0x7F, 0x7F))
            
            p_p_title = doc.add_paragraph()
            p_p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_p_title.add_run(passage.get('title', '風雨中的蒲公英'))
            set_font(run, font_name='Microsoft JhengHei', size_pt=12, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
            
            for para in passage.get('paragraphs', []):
                p_para = doc.add_paragraph()
                p_para.paragraph_format.left_indent = Pt(18)
                run = p_para.add_run(para)
                set_font(run, font_name='Microsoft JhengHei', size_pt=10.5)
                
            p_box_end = doc.add_paragraph()
            run = p_box_end.add_run("--------------------------------------------------")
            set_font(run, font_name='Microsoft JhengHei', size_pt=11, bold=True, color_rgb=RGBColor(0x7F, 0x7F, 0x7F))
            p_box_end.paragraph_format.space_after = Pt(12)
            
        # Add Question Description
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(10)
        p_q.paragraph_format.space_after = Pt(6)
        
        run_num = p_q.add_run(f"({q['qindex']}) ")
        set_font(run_num, font_name='Microsoft JhengHei', size_pt=11, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
        
        add_formatted_text(p_q, q['text'], font_name='Microsoft JhengHei', size_pt=11, bold=True)
        
        # Add Options
        for opt in q['options']:
            p_opt = doc.add_paragraph()
            p_opt.paragraph_format.left_indent = Pt(24)
            p_opt.paragraph_format.space_after = Pt(2)
            
            # Check if this option is the correct one to mark green for teachers
            is_correct_opt = False
            opt_marker_match = re.match(r'^\(([A-D])\)', opt)
            if opt_marker_match and opt_marker_match.group(1) == q['correct']:
                is_correct_opt = True
                
            run_opt = p_opt.add_run(opt)
            if is_correct_opt:
                set_font(run_opt, font_name='Microsoft JhengHei', size_pt=10.5, bold=True, color_rgb=RGBColor(0x2E, 0x7D, 0x32))
            else:
                set_font(run_opt, font_name='Microsoft JhengHei', size_pt=10.5, color_rgb=RGBColor(0x40, 0x40, 0x40))
                
        # Add Explanation
        p_exp = doc.add_paragraph()
        p_exp.paragraph_format.left_indent = Pt(24)
        p_exp.paragraph_format.space_before = Pt(4)
        p_exp.paragraph_format.space_after = Pt(14)
        
        run_ans_marker = p_exp.add_run(f"👉 正確答案：({q['correct']})\n")
        set_font(run_ans_marker, font_name='Microsoft JhengHei', size_pt=10, bold=True, color_rgb=RGBColor(0x2E, 0x7D, 0x32))
        
        # Format the explanation lines
        exp_lines = q['explanation'].split('\n')
        for exp_line in exp_lines:
            if not exp_line.strip():
                continue
            run_line = p_exp.add_run(exp_line + "\n")
            set_font(run_line, font_name='Microsoft JhengHei', size_pt=9.5, color_rgb=RGBColor(0x7F, 0x7F, 0x7F))
            
    doc.save(output_path)
    print(f"✓ Word 檔已成功生成：{output_path}")

def build_md(passage, sections, questions, output_path):
    md_lines = []
    md_lines.append("# 四年級下學期國語科期末段考複習卷（第7-12課）")
    md_lines.append("\n測驗範圍：第七課至第十二課\n")
    
    for q_idx, q in enumerate(questions):
        # Insert Section Header
        if q_idx == 0 and len(sections) > 0:
            md_lines.append(f"\n## {sections[0]}\n")
        elif q_idx == 2 and len(sections) > 1:
            md_lines.append(f"\n## {sections[1]}\n")
        elif q_idx == 4 and len(sections) > 2:
            md_lines.append(f"\n## {sections[2]}\n")
        elif q_idx == 7 and len(sections) > 3:
            md_lines.append(f"\n## {sections[3]}\n")
        elif q_idx == 12:
            if len(sections) > 4:
                md_lines.append(f"\n## {sections[4]}\n")
            
            md_lines.append("### 📖 閱讀短文")
            md_lines.append(f"#### {passage.get('title', '風雨中的蒲公英')}")
            for para in passage.get('paragraphs', []):
                md_lines.append(f"> {para}")
            md_lines.append("\n---\n")
            
        md_lines.append(f"### 第 {q['qindex']} 題")
        md_lines.append(q['text'])
        md_lines.append("")
        for opt in q['options']:
            md_lines.append(f"* {opt}")
        md_lines.append("")
        md_lines.append(f"> **【正確答案】**：({q['correct']})")
        
        # Split explanation by divs/newlines
        exp_lines = q['explanation'].split('\n')
        for exp_line in exp_lines:
            if exp_line.strip():
                md_lines.append(f"> {exp_line}")
        md_lines.append("\n---\n")
        
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_lines))
        
    print(f"✓ Markdown 檔已成功生成：{output_path}")

def main():
    html_path = 'c:/2026Antigravity0606/GEMS/output/chinese_review_exam_7_12.html'
    docx_path = 'c:/2026Antigravity0606/GEMS/output/國語_四年級_練習題.docx'
    md_path = 'c:/2026Antigravity0606/GEMS/output/國語_四年級_練習題.md'
    
    if not os.path.exists(html_path):
        print(f"Error: HTML file not found at {html_path}")
        return
        
    passage, sections, questions = parse_html_exam(html_path)
    build_docx(passage, sections, questions, docx_path)
    build_md(passage, sections, questions, md_path)

if __name__ == '__main__':
    main()
