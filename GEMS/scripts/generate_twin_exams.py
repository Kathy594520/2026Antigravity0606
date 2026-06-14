#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
新竹縣芎林國民小學 114 學年度第二學期期末定期評量四年級數學領域
雙胞胎仿生題 Word & Markdown 產出腳本
"""

import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

def set_font(run, font_name='Microsoft JhengHei', size_pt=10.5, bold=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def generate_exam_docx(docx_path):
    doc = Document()
    
    # Page Setup - Margins
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('新竹縣芎林國民小學 114 學年度第二學期期末定期評量\n四年級數學領域 雙胞胎仿生題練習卷')
    set_font(run_title, font_name='Microsoft JhengHei', size_pt=15, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
    p_title.paragraph_format.space_after = Pt(12)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run('四年____班  座號：____  姓名：_______________  家長簽名：_____________')
    set_font(run_sub, font_name='Microsoft JhengHei', size_pt=11, bold=True)
    p_sub.paragraph_format.space_after = Pt(20)

    # --- PART 1: STUDENT VERSION ---
    
    # Section 1 Header
    p_sec1 = doc.add_paragraph()
    run_sec1 = p_sec1.add_run('一、選擇題 (每題 3 分，共 30 分)')
    set_font(run_sec1, font_name='Microsoft JhengHei', size_pt=12, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
    p_sec1.paragraph_format.space_before = Pt(12)
    p_sec1.paragraph_format.space_after = Pt(6)
    
    # MC Questions
    mc_questions = [
        {
            "num": 1,
            "desc": "邊長為 200 公分的正方形，面積是多少？",
            "options": ["(1) 4 平方公分", "(2) 400 平方公分", "(3) 4 平方公尺", "(4) 40000 平方公尺"]
        },
        {
            "num": 2,
            "desc": "已經知道「125 × 70 × 8 × 5 = 1000 × 乙」，請問乙是多少？",
            "options": ["(1) 35", "(2) 350", "(3) 70", "(4) 50"]
        },
        {
            "num": 3,
            "desc": "下列哪一個算式的計算結果是正確的？",
            "options": [
                "(1) 150 - 70 + 30 = 150 - 100 = 50",
                "(2) 40 × (30 - 5) = 40 × 25 = 1000",
                "(3) 80 + 20 × 2 = 100 × 2 = 200",
                "(4) 60 + 36 ÷ 6 = 96 ÷ 6 = 16"
            ]
        },
        {
            "num": 4,
            "desc": "將 9/2、3.6、85/100、5.2 這四個數標示在數線上，哪一個數在數線的最右邊？",
            "options": ["(1) 9/2", "(2) 3.6", "(3) 85/100", "(4) 5.2"]
        },
        {
            "num": 5,
            "desc": "城市美術館上午 9 時開始開放，一直到下午 10 時閉館。請問美術館一天開放多少小時？",
            "options": ["(1) 11 小時", "(2) 12 小時", "(3) 13 小時", "(4) 14 小時"]
        },
        {
            "num": 6,
            "desc": "下列關於 1 立方公分正方體積木形體的描述，哪一個造型無法用剛好 6 個積木排出？",
            "options": [
                "(1) 由 6 個積木排成的 L 形立體形體",
                "(2) 由 6 個積木排成兩排的長方體",
                "(3) 由 6 個積木組成的三層階梯造型",
                "(4) 底層放 4 個、中層放 2 個、最上層放 1 個的造型"
            ]
        },
        {
            "num": 7,
            "desc": "下列有幾個分數的值與 2 相等？\n「 10/5 、 14/7 、 6/3 、 8/2 、 9/9 、 12/6 、 16/4 」",
            "options": ["(1) 3 個", "(2) 4 個", "(3) 5 個", "(4) 6 個"]
        },
        {
            "num": 8,
            "desc": "每個積木是 1 立方公分。甲形體由 14 個積木組成，乙形體由 12 個積木組成，請問這兩個形體的體積相差多少立方公分？",
            "options": ["(1) 1 立方公分", "(2) 2 立方公分", "(3) 3 立方公分", "(4) 4 立方公分"]
        },
        {
            "num": 9,
            "desc": "小華晚上 9 時 30 分上床睡覺，隔天早上 6 時 00 分起床，請問小華一共睡了多久時間？",
            "options": ["(1) 7 小時 30 分鐘", "(2) 8 小時 30 分鐘", "(3) 9 小時", "(4) 8 小時"]
        },
        {
            "num": 10,
            "desc": "A 是長方形，寬為 3 公分；B 是正方形，邊長為 4 公分。現在用兩個 A 垂直並排、兩個 B 垂直疊起組合成一個大長方形（如下圖所示，左側為並排的 A，右側為疊起的 B，高度剛好相等），算出這個組合圖形的周長是多少？",
            "options": ["(1) 28 公分", "(2) 36 公分", "(3) 38 公分", "(4) 42 公分"]
        }
    ]
    
    for q in mc_questions:
        p_q = doc.add_paragraph()
        run_num = p_q.add_run(f"({q['num']}) ")
        set_font(run_num, font_name='Microsoft JhengHei', size_pt=10.5, bold=True)
        run_desc = p_q.add_run(q['desc'])
        set_font(run_desc, font_name='Microsoft JhengHei', size_pt=10.5)
        p_q.paragraph_format.left_indent = Pt(12)
        p_q.paragraph_format.space_after = Pt(4)
        
        # Options in a single line or two lines
        p_opt = doc.add_paragraph()
        p_opt.paragraph_format.left_indent = Pt(28)
        p_opt.paragraph_format.space_after = Pt(8)
        for idx, opt in enumerate(q['options']):
            run_opt = p_opt.add_run(f"{opt}   ")
            set_font(run_opt, font_name='Microsoft JhengHei', size_pt=10)
            
    # Section 2 Header
    p_sec2 = doc.add_paragraph()
    run_sec2 = p_sec2.add_run('二、填填看 (每格 2 分，共 30 分)')
    set_font(run_sec2, font_name='Microsoft JhengHei', size_pt=12, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
    p_sec2.paragraph_format.space_before = Pt(16)
    p_sec2.paragraph_format.space_after = Pt(6)
    
    # Fill in the blanks
    fill_questions = [
        "1. 下面是一條以 0.1 (十分之一) 為刻度的數線。箭頭所指的刻度各是多少？用帶分數表示。\n"
        "   |----|----|----|----|----|----|----|----|----|----|----|----|----|----|----|----|----|----|\n"
        "   3    ↑                                             5         ↑                        6\n"
        "   答：左側箭頭指的刻度是 (        ) ；右側箭頭指的刻度是 (        ) 。",
        
        "2. 在括號中填入正確的數：\n"
        "   2/5 ＝ (      )/15 ＝ 12/(      ) ＝ (      )/20",
        
        "4. 時間單位換算：\n"
        "   14 分鐘 45 秒鐘 ＝ (          ) 秒鐘",
        
        "5. 時間與日數換算：\n"
        "   110 小時 ＝ (          ) 日 (          ) 小時",
        
        "6. 時間多重換算：\n"
        "   3 小時 20 分鐘 ＝ (          ) 分鐘 ＝ (          ) 秒鐘",
        
        "10. 體積保存性判斷：\n"
        "    爸爸把一個大西瓜切成 8 片，西瓜的總體積會（ 變大、變小、不改變 ）。\n"
        "    答：(          ) 。",
        
        "11. 每個小正方體積木的體積是 1 立方公分，數數看，下面各形體的體積是多少立方公分？\n"
        "    ① 底部長 3 個、寬 2 個積木，第二層長 2 個、寬 2 個積木，最上層長 1 個、寬 2 個積木拼成的階梯形體。\n"
        "       答：(          ) 立方公分。\n"
        "    ② 長 4 個、寬 3 個、高 2 個積木排成的長方體。\n"
        "       答：(          ) 立方公分。\n"
        "    ③ 底層 3×3 正方形，第二層 2×2 正方形，最上層 1 個積木疊成的金字塔形體。\n"
        "       答：(          ) 立方公分。\n"
        "    ④ 長 3 個、寬 3 個、高 2 個的長方體中，拿走最上層的 3 個積木後剩餘的形體。\n"
        "       答：(          ) 立方公分。"
    ]
    
    for item in fill_questions:
        p_f = doc.add_paragraph()
        run_f = p_f.add_run(item)
        set_font(run_f, font_name='Microsoft JhengHei', size_pt=10.5)
        p_f.paragraph_format.left_indent = Pt(12)
        p_f.paragraph_format.space_after = Pt(10)

    # Section 3 Header
    p_sec3 = doc.add_paragraph()
    run_sec3 = p_sec3.add_run('三、應用題 (寫出算式與答案) (每題 4 分，共 40 分)')
    set_font(run_sec3, font_name='Microsoft JhengHei', size_pt=12, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
    p_sec3.paragraph_format.space_before = Pt(16)
    p_sec3.paragraph_format.space_after = Pt(6)
    
    app_questions = [
        "1. 林小軒今年暑假計畫去日本東京遊玩，他搭乘的航班於上午 2 時 45 分從臺灣桃園國際機場出發，於上午 7 時 15 分抵達日本東京成田機場。請問：這趟航班的飛行時間是多少？",
        
        "2. 果園今天採收了 18 箱橘子，每箱裝有 45 顆。果農要將這些橘子平均分裝送給 9 位鄰居，請問每位鄰居可以收到多少顆橘子？\n（注意：請把問題寫成一個合併算式，並利用巧算或比較好算的方法算出答案。）",
        
        "3. 慢跑好手林小明於 2026 年 3 月 15 日上午 8 時整出發參加城市半程馬拉松賽事，他一路上保持穩定的步伐，最後耗時 85 分鐘順利抵達終點線。請問：他抵達終點時是上午幾時幾分？",
        
        "4. 媽媽每個月固定訂購 4 次親子教育雜誌，每次的雜誌運費與內容費用共是 250 元。請問：媽媽一年下來共需要支付多少雜誌訂閱費？\n（注意：請把問題寫成一個合併算式，並利用結合律或比較好算的方法算出答案。）",
        
        "5. 一箱蘋果有 24 顆。小新拿了 5/6 箱蘋果，小葵拿了 17/24 箱蘋果。請問：誰拿的蘋果比較多？多幾分之幾箱？",
        
        "6. 社區活動廣場準備鋪設人工綠草皮（如圖所示，外部是一個邊長 15 公尺的大正方形，中間挖空了一個邊長 6 公尺的正方形噴水池，塗色部分為草皮區），請問這片草皮區的面積是多少平方公尺？",
        
        "7. 數線上的每一個刻度代表 1/5 公分。有一隻蝸牛停在數線上 3又4/5 公分的位置，如果牠往左爬行了 3/5 公分，請問蝸牛會停在數線上什麼位置？（請用帶分數表示答案）",
        
        "8. 有一個正方形的周長是 120 公分，請問這個正方形的面積是多少平方公分？",
        
        "9. 小明帶著 1500 元去文具行購買開學用品，他買了一本 264 元的英漢字典和一盒 536 元的專業畫冊，請問小明結帳後還剩下多少元？\n（注意：請把問題寫成一個合併算式，並利用巧算或湊整數的方式算出答案。）",
        
        "10. 丙是一個正方體圖卡的面（即正方形）。如下圖所示，是由四個完全相同的正方形「丙」橫向並排貼在一起所組合成的長方形。若已知正方形丙的邊長是 4 公分，請問這個組合長方形的周長是多少公分？"
    ]
    
    for idx, item in enumerate(app_questions):
        p_a = doc.add_paragraph()
        run_idx = p_a.add_run(f"{idx+1}. ")
        set_font(run_idx, font_name='Microsoft JhengHei', size_pt=10.5, bold=True)
        run_a = p_a.add_run(item)
        set_font(run_a, font_name='Microsoft JhengHei', size_pt=10.5)
        p_a.paragraph_format.left_indent = Pt(12)
        p_a.paragraph_format.space_after = Pt(40)  # Leave blank space for student writing

    # --- PART 2: ANSWER KEY AND EXPLANATIONS ---
    doc.add_page_break()
    
    p_ans_title = doc.add_paragraph()
    p_ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ans_title = p_ans_title.add_run('新竹縣芎林國民小學 114 學年度第二學期期末定期評量\n四年級數學領域 雙胞胎仿生題參考答案與詳解')
    set_font(run_ans_title, font_name='Microsoft JhengHei', size_pt=14, bold=True, color_rgb=RGBColor(0x2E, 0x7D, 0x32))
    p_ans_title.paragraph_format.space_after = Pt(20)

    # Section 1 Answers
    p_sec1_a = doc.add_paragraph()
    run_sec1_a = p_sec1_a.add_run('一、選擇題答案與詳解')
    set_font(run_sec1_a, font_name='Microsoft JhengHei', size_pt=12, bold=True, color_rgb=RGBColor(0x2E, 0x7D, 0x32))
    
    mc_answers = [
        ("3", "200公分 = 2公尺。正方形面積 = 邊長 × 邊長 = 2公尺 × 2公尺 = 4平方公尺。故選(3)。"),
        ("2", "利用乘法結合律，先算「125 × 8 = 1000」與「70 × 5 = 350」，算式化為 1000 × 350。因此乙 = 350。故選(2)。"),
        ("2", "檢查選項：(1) 應先算減法再算加法，150-70+30=80+30=110；(2) 括號先算，40 × 25 = 1000，正確；(3) 先乘後加，80+40=120；(4) 先除後加，60+9=69。故選(2)。"),
        ("4", "將分數換算為小數：9/2 = 4.5，85/100 = 0.85。比較四數大小：0.85 < 3.6 < 4.5 (9/2) < 5.2。最右邊即最大值，為 5.2。故選(4)。"),
        ("3", "下午 10 時是 22 時。開放時間 = 結束時刻 - 開始時刻 = 22時 - 9時 = 13小時。故選(3)。"),
        ("4", "(1)、(2)、(3)選項積木數量皆為 6 個；(4)選項積木數量為 4 + 2 + 1 = 7 個，無法用 6 個排出。故選(4)。"),
        ("2", "值與 2 相等的分數分子應是分母的 2 倍：10/5=2、14/7=2、6/3=2、12/6=2，共 4 個。其餘：8/2=4、9/9=1、16/4=4。故選(2)。"),
        ("2", "甲體積為 14 立方公分，乙體積為 12 立方公分。體積差 = 14 - 12 = 2 立方公分。故選(2)。"),
        ("2", "從晚上 9 時 30 分到午夜 12 時 (24時) 是 2 小時 30 分鐘；從午夜到早上 6 時是 6 小時。一共睡了 8 小時 30 分鐘。故選(2)。"),
        ("2", "B疊兩個高度為 4 + 4 = 8 公分，代表長方形 A 的長是 8 公分。組合後的大長方形寬為 A的寬×2 加上 B的邊長 ＝ 3 + 3 + 4 = 10 公分。周長 ＝ (長 + 寬) × 2 ＝ (8 + 10) × 2 ＝ 36 公分。故選(2)。")
    ]
    
    for idx, (ans, detail) in enumerate(mc_answers):
        p_ans = doc.add_paragraph()
        run_q = p_ans.add_run(f"第 {idx+1} 題：")
        set_font(run_q, font_name='Microsoft JhengHei', size_pt=10, bold=True)
        run_ans = p_ans.add_run(f"({ans})   ")
        set_font(run_ans, font_name='Microsoft JhengHei', size_pt=10, bold=True, color_rgb=RGBColor(0x2E, 0x7D, 0x32))
        run_det = p_ans.add_run(detail)
        set_font(run_det, font_name='Microsoft JhengHei', size_pt=9.5, color_rgb=RGBColor(0x40, 0x40, 0x40))
        p_ans.paragraph_format.left_indent = Pt(12)
        p_ans.paragraph_format.space_after = Pt(4)

    # Section 2 Answers
    p_sec2_a = doc.add_paragraph()
    run_sec2_a = p_sec2_a.add_run('\n二、填填看答案與詳解')
    set_font(run_sec2_a, font_name='Microsoft JhengHei', size_pt=12, bold=True, color_rgb=RGBColor(0x2E, 0x7D, 0x32))
    p_sec2_a.paragraph_format.space_before = Pt(12)
    
    fill_answers = [
        "1. 刻度為十分之一 (0.1)。左側箭頭在 3 往右 1 格，為 3又10分之1 (或 3.1)；右側箭頭在 5 往右 7 格，為 5又10分之7 (或 5.7)。",
        "2. 2/5 ＝ 6/15 (同乘3) ＝ 12/30 (同乘6) ＝ 8/20 (同乘4)。答案為：6、30、8。",
        "4. 14 分鐘 ＝ 14 × 60 ＝ 840 秒，840 ＋ 45 ＝ 885 秒。答案為：885。",
        "5. 110 ÷ 24 ＝ 4 餘 14。答案為：4 日 14 小時。",
        "6. 3 小時 20 分鐘 ＝ 3 × 60 ＋ 20 ＝ 200 分鐘。200 × 60 ＝ 12000 秒。答案為：200、12000。",
        "10. 固體物質切開後形狀改變但所占空間總量不變。答案為：不改變。",
        "11. 積木體積計數：\n"
        "    ① 第一層有 1×2＝2 個，第二層有 2×2＝4 個，第三層有 3×2＝6 個，總共 2＋4＋6＝12 立方公分。\n"
        "    ② 長方體積木總數 ＝ 長 4 × 寬 3 × 高 2 ＝ 24 立方公分。\n"
        "    ③ 第一層有 1 個，第二層有 2×2＝4 個，第三層有 3×3＝9 個，總共 1＋4＋9＝14 立方公分。\n"
        "    ④ 完整體積為 3 × 3 × 2 ＝ 18。拿走上層 3 個剩餘 18 - 3 ＝ 15 立方公分。"
    ]
    
    for item in fill_answers:
        p_ans = doc.add_paragraph()
        run_f = p_ans.add_run(item)
        set_font(run_f, font_name='Microsoft JhengHei', size_pt=9.5, color_rgb=RGBColor(0x40, 0x40, 0x40))
        p_ans.paragraph_format.left_indent = Pt(12)
        p_ans.paragraph_format.space_after = Pt(4)

    # Section 3 Answers
    p_sec3_a = doc.add_paragraph()
    run_sec3_a = p_sec3_a.add_run('\n三、應用題答案與詳解')
    set_font(run_sec3_a, font_name='Microsoft JhengHei', size_pt=12, bold=True, color_rgb=RGBColor(0x2E, 0x7D, 0x32))
    p_sec3_a.paragraph_format.space_before = Pt(12)
    
    app_answers = [
        "1. 算式：7時15分 － 2時45分 ＝ 4小時30分鐘。\n   解析：15分不夠減45分，向小時借1得60分，60+15-45=30分，小時剩6，6-2=4小時。答：4小時30分鐘。",
        "2. 算式：45 × 18 ÷ 9 ＝ 45 × (18 ÷ 9) ＝ 45 × 2 ＝ 90。\n   解析：先將箱數除以鄰居人數，可以得到每人分得 2 箱，每箱 45 顆所以共得 90 顆。答：90顆。",
        "3. 算式：上午8時 ＋ 85分鐘 ＝ 9時25分。\n   解析：85分鐘 ＝ 1小時25分鐘。上午 8時 + 1小時25分 = 上午 9時25分。答：上午9時25分。",
        "4. 算式：250 × 4 × 12 ＝ (250 × 4) × 12 ＝ 1000 × 12 ＝ 12000。\n   解析：一月訂4次，一年有12個月。利用乘法結合律先算 250×4=1000，再乘以12。答：12000元。",
        "5. 算式：5/6 ＝ 20/24，20/24 ＞ 17/24。相差：20/24 － 17/24 ＝ 3/24 ＝ 1/8。\n   解析：小新拿了 20/24 箱，小葵拿了 17/24 箱，小新較多。答：小新比較多，多 1/8 箱 (或 3/24 箱)。",
        "6. 算式：15 × 15 － 6 × 6 ＝ 225 － 36 ＝ 189。\n   解析：大正方形面積減去中間小正方形噴水池面積。答：189平方公尺。",
        "7. 算式：3又4/5 － 3/5 ＝ 3又1/5。\n   解析：往左爬行為減法，3 4/5 - 3/5 = 3 1/5 公分。答：3又5分之1公分。",
        "8. 算式：120 ÷ 4 ＝ 30，30 × 30 ＝ 900。\n   解析：正方形邊長 ＝ 周長 ÷ 4 = 30公分。面積 = 邊長 × 邊長 = 900。答：900平方公分。",
        "9. 算式：1500 － 264 － 536 ＝ 1500 － (264 ＋ 536) ＝ 1500 － 800 ＝ 700。\n   解析：連減兩個數等於減去這兩個數的和，先將兩本書的價格相加得到 800 元，再由 1500 減去。答：700元。",
        "10. 算式：組合長方形長 ＝ 4 × 4 ＝ 16公分，寬 ＝ 4公分。周長 ＝ (16 ＋ 4) × 2 ＝ 40公分。\n    解析：四個正方形橫向貼起，長為4個邊長，寬為1個邊長。答：40公分。"
    ]
    
    for idx, item in enumerate(app_answers):
        p_ans = doc.add_paragraph()
        run_idx = p_ans.add_run(f"{idx+1}. ")
        set_font(run_idx, font_name='Microsoft JhengHei', size_pt=10, bold=True)
        run_f = p_ans.add_run(item)
        set_font(run_f, font_name='Microsoft JhengHei', size_pt=9.5, color_rgb=RGBColor(0x40, 0x40, 0x40))
        p_ans.paragraph_format.left_indent = Pt(12)
        p_ans.paragraph_format.space_after = Pt(6)

    doc.save(docx_path)
    print(f"✓ Word 檔已成功生成：{docx_path}")

def generate_exam_md(md_path):
    md_lines = []
    md_lines.append("# 新竹縣芎林國民小學 114 學年度第二學期期末定期評量")
    md_lines.append("## 四年級數學領域 雙胞胎仿生題練習卷")
    md_lines.append("\n班級：____  座號：____  姓名：_______________\n")
    
    md_lines.append("## 一、選擇題 (每題 3 分，共 30 分)")
    mc_questions = [
        ("1", "邊長為 200 公分的正方形，面積是多少？\n* (1) 4 平方公分\n* (2) 400 平方公分\n* (3) 4 平方公尺\n* (4) 40000 平方公尺"),
        ("2", "已經知道「125 × 70 × 8 × 5 = 1000 × 乙」，請問乙是多少？\n* (1) 35\n* (2) 350\n* (3) 70\n* (4) 50"),
        ("3", "下列哪一個算式的計算結果是正確的？\n* (1) 150 - 70 + 30 = 150 - 100 = 50\n* (2) 40 × (30 - 5) = 40 × 25 = 1000\n* (3) 80 + 20 × 2 = 100 × 2 = 200\n* (4) 60 + 36 ÷ 6 = 96 ÷ 6 = 16"),
        ("4", "將 9/2、3.6、85/100、5.2 這四個數標示在數線上，哪一個數在數線的最右邊？\n* (1) 9/2\n* (2) 3.6\n* (3) 85/100\n* (4) 5.2"),
        ("5", "城市美術館上午 9 時開始開放，一直到下午 10 時閉館。請問美術館一天開放多少小時？\n* (1) 11 小時\n* (2) 12 小時\n* (3) 13 小時\n* (4) 14 小時"),
        ("6", "下列關於 1 立方公分正方體積木形體的描述，哪一個造型無法用剛好 6 個積木排出？\n* (1) 由 6 個積木排成的 L 形立體形體\n* (2) 由 6 個積木排成兩排的長方體\n* (3) 由 6 個積木組成的三層階梯造型\n* (4) 底層放 4 個、中層放 2 個、最上層放 1 個的造型"),
        ("7", "下列有幾個分數的值與 2 相等？\n「 10/5 、 14/7 、 6/3 、 8/2 、 9/9 、 12/6 、 16/4 」\n* (1) 3 個\n* (2) 4 個\n* (3) 5 個\n* (4) 6 個"),
        ("8", "每個積木是 1 立方公分。甲形體由 14 個積木組成，乙形體由 12 個積木組成，請問這兩個形體的體積相差多少立方公分？\n* (1) 1 立方公分\n* (2) 2 立方公分\n* (3) 3 立方公分\n* (4) 4 立方公分"),
        ("9", "小華晚上 9 時 30 分上床睡覺，隔天早上 6 時 00 分起床，請問小華一共睡了多久時間？\n* (1) 7 小時 30 分鐘\n* (2) 8 小時 30 分鐘\n* (3) 9 小時\n* (4) 8 小時"),
        ("10", "A 是長方形，寬為 3 公分；B 是正方形，邊長為 4 公分。現在用兩個 A 垂直並排、兩個 B 垂直疊起組合成一個大長方形（左側為並排的 A，右側為疊起的 B，高度剛好相等），算出這個組合圖形的周長是多少？\n* (1) 28 公分\n* (2) 36 公分\n* (3) 38 公分\n* (4) 42 公分")
    ]
    
    for num, q in mc_questions:
        md_lines.append(f"### 第 {num} 題")
        md_lines.append(q)
        md_lines.append("")
        
    md_lines.append("\n## 二、填填看 (每格 2 分，共 30 分)")
    fill_questions = [
        "1. 下面是一條以 0.1 (十分之一) 為刻度的數線。箭頭所指的刻度各是多少？用帶分數表示。\n"
        "   `|----|----|----|----|----|----|----|----|----|----|----|----|----|----|----|----|----|----|`\n"
        "   `3    ↑                                             5         ↑                        6`\n"
        "   * 左側箭頭指的刻度是：_____\n"
        "   * 右側箭頭指的刻度是：_____",
        
        "2. 在括號中填入正確的數：\n"
        "   `2/5 ＝ (  a  )/15 ＝ 12/(  b  ) ＝ (  c  )/20`\n"
        "   * a = _____, b = _____, c = _____",
        
        "4. 時間單位換算：\n"
        "   `14 分鐘 45 秒鐘 ＝ (      ) 秒鐘`",
        
        "5. 時間與日數換算：\n"
        "   `110 小時 ＝ (      ) 日 (      ) 小時`",
        
        "6. 時間多重換算：\n"
        "   `3 小時 20 分鐘 ＝ (      ) 分鐘 ＝ (      ) 秒鐘`",
        
        "10. 體積保存性判斷：\n"
        "    爸爸把一個大西瓜切成 8 片，西瓜的總體積會（ 變大、變小、不改變 ）。\n"
        "    答：(          )",
        
        "11. 每個小正方體積木的體積是 1 立方公分，數數看，下面各形體的體積是多少立方公分？\n"
        "    * ① 底部長 3 個、寬 2 個積木，第二層長 2 個、寬 2 個積木，最上層長 1 個、寬 2 個拼成的階梯形體：(      ) 立方公分\n"
        "    * ② 長 4 個、寬 3 個、高 2 個積木排成的長方體：(      ) 立方公分\n"
        "    * ③ 底層 3×3，第二層 2×2，最上層 1 個疊成的金字塔形體：(      ) 立方公分\n"
        "    * ④ 長 3 個、寬 3 個、高 2 個長方體中，拿走最上層的 3 個積木後剩餘的形體：(      ) 立方公分"
    ]
    
    for q in fill_questions:
        md_lines.append(q)
        md_lines.append("")
        
    md_lines.append("\n## 三、應用題 (寫出算式與答案) (每題 4 分，共 40 分)")
    app_questions = [
        "1. 林小軒今年暑假計畫去日本東京遊玩，他搭乘的航班於上午 2 時 45 分從臺灣桃園國際機場出發，於上午 7 時 15 分抵達日本東京成田機場。請問：這趟航班的飛行時間是多少？",
        "2. 果園今天採收了 18 箱橘子，每箱裝有 45 顆。果農要將這些橘子平均分裝送給 9 位鄰居，請問每位鄰居可以收到多少顆橘子？\n（注意：請把問題寫成一個合併算式，並利用巧算或比較好算的方法算出答案。）",
        "3. 慢跑好手林小明於 2026 年 3 月 15 日上午 8 時整出發參加城市半程馬拉松賽事，他一路上保持穩定的步伐，最後耗時 85 分鐘順利抵達終點線。請問：他抵達終點時是上午幾時幾分？",
        "4. 媽媽每個月固定訂購 4 次親子教育雜誌，每次的雜誌運費與內容費用共是 250 元。請問：媽媽一年下來共需要支付多少雜誌訂閱費？\n（注意：請把問題寫成一個合併算式，並利用結合律或比較好算的方法算出答案。）",
        "5. 一箱蘋果有 24 顆。小新拿了 5/6 箱蘋果，小葵拿了 17/24 箱蘋果。請問：誰拿的蘋果比較多？多幾分之幾箱？",
        "6. 社區活動廣場準備鋪設人工綠草皮（外部是一個邊長 15 公尺的大正方形，中間挖空了一個邊長 6 公尺的正方形噴水池，塗色部分為草皮區），請問這片草皮區的面積是多少平方公尺？",
        "7. 數線上的每一個刻度代表 1/5 公分。有一隻蝸牛停在數線上 3又4/5 公分的位置，如果牠往左爬行了 3/5 公分，請問蝸牛會停在數線上什麼位置？（請用帶分數表示答案）",
        "8. 有一個正方形的周長是 120 公分，請問這個正方形的面積是多少平方公分？",
        "9. 小明帶著 1500 元去文具行購買開學用品，他買了一本 264 元的英漢字典和一盒 536 元的專業畫冊，請問小明結帳後還剩下多少元？\n（注意：請把問題寫成一個合併算式，並利用巧算或湊整數的方式算出答案。）",
        "10. 丙是一個正方體圖卡的面（即正方形）。是由四個完全相同的正方形「丙」橫向並排貼在一起所組合成的長方形。若已知正方形丙的邊長是 4 公分，請問這個組合長方形的周長是多少公分？"
    ]
    
    for idx, q in enumerate(app_questions):
        md_lines.append(f"### 第 {idx+1} 題")
        md_lines.append(q)
        md_lines.append("")
        
    md_lines.append("\n---\n")
    md_lines.append("## 雙胞胎仿生題 參考答案與詳解")
    
    md_lines.append("\n### 一、選擇題答案")
    mc_answers = [
        ("1", "(3)", "200公分 = 2公尺。面積 = 2 × 2 = 4平方公尺。"),
        ("2", "(2)", "125 × 70 × 8 × 5 = (125 × 8) × (70 × 5) = 1000 × 350。"),
        ("3", "(2)", "40 × (30 - 5) = 40 × 25 = 1000 正確。"),
        ("4", "(4)", "9/2=4.5，85/100=0.85，數線最右邊為最大值 5.2。"),
        ("5", "(3)", "下午 10 時 = 22 時。22 - 9 = 13 小時。"),
        ("6", "(4)", "(4)選項底層 4 個、中層 2 個、上層 1 個共需 7 個積木。"),
        ("7", "(2)", "值為 2 的有 10/5、14/7、6/3、12/6 共 4 個。"),
        ("8", "(2)", "體積差 = 14 - 12 = 2 立方公分。"),
        ("9", "(2)", "晚上 9:30 至 12:00 是 2.5 小時，加早上 6 小時共 8 小時 30 分鐘。"),
        ("10", "(2)", "組合長方形長 8 公分，寬 3+3+4=10 公分。周長 = (8+10)×2 = 36 公分。")
    ]
    for num, ans, explanation in mc_answers:
        md_lines.append(f"* **第 {num} 題**：答案 **{ans}**。解析：{explanation}")
        
    md_lines.append("\n### 二、填填看答案")
    fill_answers = [
        "* **1.** 左側 `3又10分之1` (或 `3.1`)；右側 `5又10分之7` (或 `5.7`)",
        "* **2.** `6`、`30`、`8`",
        "* **4.** `885` 秒",
        "* **5.** `4` 日 `14` 小時",
        "* **6.** `200` 分鐘、`12000` 秒",
        "* **10.** `不改變`",
        "* **11.** ① `12` 立方公分；② `24` 立方公分；③ `14` 立方公分；④ `15` 立方公分"
    ]
    for ans in fill_answers:
        md_lines.append(ans)
        
    md_lines.append("\n### 三、應用題答案")
    app_answers = [
        "1. **算式**：7時15分 － 2時45分 ＝ 4小時30分鐘。\n   **答**：4小時30分鐘。",
        "2. **算式**：45 × 18 ÷ 9 ＝ 45 × (18 ÷ 9) ＝ 45 × 2 ＝ 90。\n   **答**：90顆。",
        "3. **算式**：上午8時 ＋ 85分鐘 ＝ 上午9時25分。\n   **答**：上午9時25分。",
        "4. **算式**：250 × 4 × 12 ＝ (250 × 4) × 12 ＝ 1000 × 12 ＝ 12000。\n   **答**：12000元。",
        "5. **算式**：5/6 ＝ 20/24，20/24 ＞ 17/24。相差：20/24 － 17/24 ＝ 3/24 ＝ 1/8。\n   **答**：小新比較多，多 1/8 箱 (或 3/24 箱)。",
        "6. **算式**：15 × 15 － 6 × 6 ＝ 225 － 36 ＝ 189。\n   **答**：189平方公尺。",
        "7. **算式**：3又4/5 － 3/5 ＝ 3又1/5。\n   **答**：3又5分之1公分。",
        "8. **算式**：120 ÷ 4 ＝ 30，30 × 30 ＝ 900。\n   **答**：900平方公分。",
        "9. **算式**：1500 － 264 － 536 ＝ 1500 － (264 ＋ 536) ＝ 1500 － 800 ＝ 700。\n   **答**：700元。",
        "10. **算式**：組合長方形長 4×4=16公分，寬 4公分。周長 ＝ (16 ＋ 4) × 2 ＝ 40公分。\n    **答**：40公分。"
    ]
    for idx, ans in enumerate(app_answers):
        md_lines.append(f"**第 {idx+1} 題**：\n{ans}\n")
        
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_lines))
    print(f"✓ Markdown 檔已成功生成：{md_path}")

def main():
    docx_path = 'c:/2026Antigravity0606/GEMS/output/四年級數學_期末定期評量_雙胞胎仿生題.docx'
    md_path = 'c:/2026Antigravity0606/GEMS/output/四年級數學_期末定期評量_雙胞胎仿生題.md'
    
    generate_exam_docx(docx_path)
    generate_exam_md(md_path)

if __name__ == '__main__':
    main()
