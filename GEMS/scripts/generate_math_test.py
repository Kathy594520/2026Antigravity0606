#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
學力測驗雙胞胎練習題生成腳本 (測試用)
"""

import sys
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


def generate_docx():
    doc = Document()
    
    # 設定中文字型與大小
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # 標題
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('114 年學力檢測 數學五年級 雙胞胎練習題')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    
    # 說明文字
    desc = doc.add_paragraph()
    desc.add_run('本練習卷根據 114 年學力檢測五年級數學科試題進行「雙胞胎仿寫」，提供相同解題邏輯的仿題供學生練習。')
    desc.paragraph_format.space_after = Pt(20)
    
    # 題目清單
    questions = [
        {
            "num": 1,
            "desc": "已知 24 的倍數有 24、48、72、96、120、……。下列何者是 24 和 36 的最小公倍數？",
            "options": ["(1) 12", "(2) 48", "(3) 72", "(4) 144"],
            "ans": "(3) 72",
            "analysis": "24 的倍數有：24, 48, 72, 96, 120...；36 的倍數有：36, 72, 108...。兩者最小公倍數為 72。"
        },
        {
            "num": 2,
            "desc": "請問「2 ÷ 9」商算到小數點後第三位時的餘數為何？",
            "options": ["(1) 0.002", "(2) 0.02", "(3) 0.2", "(4) 0.0002"],
            "ans": "(1) 0.002",
            "analysis": "2 ÷ 9 = 0.222 ... 餘數。餘數 = 被除數 - 除數 × 商 = 2 - 9 × 0.222 = 2 - 1.998 = 0.002。"
        },
        {
            "num": 3,
            "desc": "將下列三個算式合併成一個算式：\n340 + 260 = 600\n600 × 4 = 2400\n3000 - 2400 = 600\n下列哪個選項正確？",
            "options": [
                "(1) 3000 - 340 + 260 × 4",
                "(2) 3000 - (340 + 260) × 4",
                "(3) (340 + 260) × 4 - 3000",
                "(4) 340 + 260 × 4 - 3000"
            ],
            "ans": "(2) 3000 - (340 + 260) × 4",
            "analysis": "先算括號內的 340+260 得到 600，再乘以 4 得到 2400，最後由 3000 減去該值。故正確算式為 3000 - (340 + 260) × 4。"
        },
        {
            "num": 4,
            "desc": "「橡皮擦一盒賣 45 元，妹妹早上買了 19 盒，下午又買了 1 盒，共花了多少元？」下列哪個算式無法算出正確答案？",
            "options": [
                "(1) 45 × (19 + 1)",
                "(2) 45 × 19 + 45",
                "(3) 45 × 19 + 19",
                "(4) 45 × 1 + 45 × 19"
            ],
            "ans": "(3) 45 × 19 + 19",
            "analysis": "正確算法應為：45 × (19 + 1) 或 45 × 19 + 45 × 1。選項 (3) 的「+ 19」代表加了 19 元而非一盒的價格，故無法算出正確答案。"
        },
        {
            "num": 5,
            "desc": "有一個正方體木箱，其中一個面的面積是 9 平方公寸，請問木箱的體積是多少立方公寸？",
            "options": ["(1) 9", "(2) 27", "(3) 81", "(4) 54"],
            "ans": "(2) 27",
            "analysis": "正方體的一個面是正方形，面積為 9 平方公寸，代表其邊長為 3 公寸 (3 × 3 = 9)。正方體的體積 = 邊長 × 邊長 × 邊長 = 3 × 3 × 3 = 27 立方公寸。"
        }
    ]
    
    # 寫入題目
    for q in questions:
        p_q = doc.add_paragraph()
        p_q.add_run(f"{q['num']}. {q['desc']}").bold = True
        p_q.paragraph_format.space_before = Pt(10)
        
        # 選項
        for opt in q['options']:
            p_opt = doc.add_paragraph(f"    {opt}")
            p_opt.paragraph_format.space_after = Pt(2)
            
        p_ans = doc.add_paragraph()
        p_ans.add_run(f"    👉 【答案】：{q['ans']}").font.color.rgb = RGBColor(0x00, 0xaa, 0x00)
        p_ans.add_run(f"\n    💡 【解析】：{q['analysis']}").font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p_ans.paragraph_format.space_after = Pt(15)
        
    doc.save('c:/2026Antigravity0606/GEMS/output/五年級數學學力測驗_雙胞胎練習題.docx')
    print("✓ Word 檔已成功生成：c:/2026Antigravity0606/GEMS/output/五年級數學學力測驗_雙胞胎練習題.docx")


def generate_md():
    md_content = """# 114 年學力檢測 數學五年級 雙胞胎練習題

本練習卷根據 114 年學力檢測五年級數學科試題進行「雙胞胎仿寫」，提供相同解題邏輯的仿題供學生練習。

---

### 1. 最小公倍數練習
已知 24 的倍數有 24、48、72、96、120、……。下列何者是 24 和 36 的最小公倍數？
* (1) 12
* (2) 48
* (3) 72
* (4) 144

> **【答案】**：(3) 72
> **【解析】**：24 的倍數有：24, 48, 72, 96, 120...；36 的倍數有：36, 72, 108...。兩者最小公倍數為 72。

---

### 2. 小數除法餘數練習
請問「2 ÷ 9」商算到小數點後第三位時的餘數為何？
* (1) 0.002
* (2) 0.02
* (3) 0.2
* (4) 0.0002

> **【答案】**：(1) 0.002
> **【解析】**：2 ÷ 9 = 0.222 ... 餘數。餘數 = 被除數 - 除數 × 商 = 2 - 9 × 0.222 = 2 - 1.998 = 0.002。

---

### 3. 算式合併練習
將下列三個算式合併成一個算式：
* 340 + 260 = 600
* 600 × 4 = 2400
* 3000 - 2400 = 600

下列哪個選項正確？
* (1) 3000 - 340 + 260 × 4
* (2) 3000 - (340 + 260) × 4
* (3) (340 + 260) × 4 - 3000
* (4) 340 + 260 × 4 - 3000

> **【答案】**：(2) 3000 - (340 + 260) × 4
> **【解析】**：先算括號內的 340+260 得到 600，再乘以 4 得到 2400，最後由 3000 減去該值。故正確算式為 3000 - (340 + 260) × 4。

---

### 4. 運算規律練習
「橡皮擦一盒賣 45 元，妹妹早上買了 19 盒，下午又買了 1 盒，共花了多少元？」下列哪個算式無法算出正確答案？
* (1) 45 × (19 + 1)
* (2) 45 × 19 + 45
* (3) 45 × 19 + 19
* (4) 45 × 1 + 45 × 19

> **【答案】**：(3) 45 × 19 + 19
> **【解析】**：正確算法應為：45 × (19 + 1) 或 45 × 19 + 45 × 1。選項 (3) 的「+ 19」代表加了 19 元而非一盒的價格，故無法算出正確答案。

---

### 5. 正方體體積與面積關係
有一個正方體木箱，其中一個面的面積是 9 平方公寸，請問木箱的體積是多少立方公寸？
* (1) 9
* (2) 27
* (3) 81
* (4) 54

> **【答案】**：(2) 27
> **【解析】**：正方體的一個面是正方形，面積為 9 平方公寸，代表其邊長為 3 公寸 (3 × 3 = 9)。正方體的體積 = 邊長 × 邊長 × 邊長 = 3 × 3 × 3 = 27 立方公寸。
"""
    with open('c:/2026Antigravity0606/GEMS/output/五年級數學學力測驗_雙胞胎練習題.md', 'w', encoding='utf-8') as f:
        f.write(md_content)
    print("✓ Markdown 檔已成功生成：c:/2026Antigravity0606/GEMS/output/五年級數學學力測驗_雙胞胎練習題.md")


if __name__ == "__main__":
    generate_docx()
    generate_md()
